#!/usr/bin/env python3

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import textwrap

def create_docx():
    doc = Document()
    
    # Remove default paragraph spacing
    style = doc.styles['Normal']
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.add_run('Национальный исследовательский университет ИТМО').bold = True
    
    doc.add_paragraph()
    center_para = doc.add_paragraph()
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_para.add_run('Мегафакультет компьютерных технологий и управления')
    
    doc.add_paragraph()
    center_para = doc.add_paragraph()
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_para.add_run('Факультет программной инженерии и компьютерной техники')
    
    # Add some space
    for _ in range(5):
        doc.add_paragraph()
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('Отчет по лабораторной работе №3')
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    doc.add_paragraph()
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run('по курсу «Языки программирования»')
    subtitle_run.font.size = Pt(14)
    
    # Add space
    for _ in range(6):
        doc.add_paragraph()
    
    # Author info
    author_para = doc.add_paragraph()
    author_para.add_run('Выполнил:').bold = True
    doc.add_paragraph('Студент группы P4119')
    doc.add_paragraph('М.С. Маршалов')
    
    for _ in range(2):
        doc.add_paragraph()
    
    teacher_para = doc.add_paragraph()
    teacher_para.add_run('Преподаватель:').bold = True
    doc.add_paragraph('Ю. Д. Кореньков')
    
    # Add space
    for _ in range(6):
        doc.add_paragraph()
    
    city_para = doc.add_paragraph()
    city_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city_para.add_run('Санкт-Петербург')
    
    doc.add_paragraph()
    year_para = doc.add_paragraph()
    year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year_para.add_run('2026')
    
    # Page break
    doc.add_page_break()
    
    # Main content sections
    content = [
        {
            'title': 'Цель',
            'content': 'Разработать компилятор с языка программирования Simple в ассемблер для различных архитектур (x86-64 Linux, x86-64 Windows, RISC-V). Реализовать генерацию кода на основе графа потока управления, поддерживающий основные конструкции языка программирования.'
        },
        {
            'title': 'Задачи',
            'content': None,
            'subsections': [
                {
                    'title': '1. Реализовать парсер языка Simple',
                    'subtasks': [
                        '1.1. Разработать синтаксический анализатор для языка Simple',
                        '1.2. Построить абстрактное синтаксическое дерево (AST)',
                        '1.3. Реализовать семантический анализ и проверку типов'
                    ]
                },
                {
                    'title': '2. Создать систему генерации кода',
                    'subtasks': [
                        '2.1. Реализовать построение графа потока управления (CFG)',
                        '2.2. Разработать генераторы ассемблерного кода для x86-64 Linux',
                        '2.3. Реализовать генерацию кода для x86-64 Windows',
                        '2.4. Создать генератор для архитектуры RISC-V'
                    ]
                },
                {
                    'title': '3. Поддержать основные конструкции языка',
                    'subtasks': [
                        '3.1. Арифметические и логические операции',
                        '3.2. Условные операторы (if-else if-else)',
                        '3.3. Циклы (while, do-while, for)',
                        '3.4. Функции и вызовы встроенных функций'
                    ]
                },
                {
                    'title': '4. Обеспечить кросс-платформенность',
                    'subtasks': [
                        '4.1. Реализовать абстракцию над целевыми платформами',
                        '4.2. Обеспечить совместимость с системными вызовами',
                        '4.3. Поддержать различные конвенции вызовов'
                    ]
                }
            ]
        },
        {
            'title': 'Ход работы',
            'content': None,
            'subsections': [
                {
                    'title': '1. Разработка парсера языка Simple',
                    'content': 'Был реализован лексический и синтаксический анализатор для языка Simple. Язык поддерживает объявления переменных, функции, управляющие конструкции и выражения.',
                    'code': '''function fibonacci() -> int {
    variable a -> int;
    variable b -> int;
    variable temp -> int;
    variable n -> int;
    
    a = 0;
    b = 1;
    n = 10;
    
    while (n > 0) {
        printf("%d\\n", a);
        temp = a + b;
        a = b;
        b = temp;
        n = n - 1;
    }
    
    return a;
}

function main() -> int {
    return fibonacci();
}'''
                },
                {
                    'title': '2. Генерация графа потока управления',
                    'content': 'На основе AST строится CFG, состоящий из базовых блоков. Каждый базовый блок содержит последовательность операций и переходы к другим блокам.',
                    'code': '''@dataclass
class BasicBlock:
    id: int
    operations: List[Operation]
    true_branch: Optional['BasicBlock']
    false_branch: Optional['BasicBlock'] 
    next_block: Optional['BasicBlock']
    is_loop_start: bool = False
    is_loop_end: bool = False'''
                },
                {
                    'title': '3. Генераторы ассемблерного кода',
                    'content': 'Реализованы три генератора кода для различных архитектур.',
                    'subtasks': [
                        '3.1. Генератор для x86-64 Linux - использует синтаксис NASM и системные вызовы Linux',
                        '3.2. Генератор для x86-64 Windows - использует синтаксис MASM и Windows API',
                        '3.3. Генератор для RISC-V - генерирует код в формате GCC для RISC-V'
                    ]
                },
                {
                    'title': '4. Тестирование и результаты',
                    'content': 'Компилятор был протестирован на нескольких примерах: калькулятор, числа Фибоначчи, сортировка массива. Все примеры успешно компилируются и выполняются на всех трех поддерживаемых платформах.'
                },
                {
                    'title': '5. Сборка и использование',
                    'content': 'Для сборки проекта используется Makefile. Основные команды:\n• make all - сборка всех примеров\n• make calculator-linux - сборка калькулятора для Linux\n• make calculator-riscv - сборка для RISC-V\n• make clean - очистка выходных файлов'
                }
            ]
        },
        {
            'title': 'Выводы',
            'content': 'В ходе лабораторной работы был успешно разработан компилятор с языка Simple в ассемблер для трех архитектур. Основные достижения:\n\n• Реализован полный цикл компиляции: от исходного кода до исполняемого файла\n• Поддерживаются основные конструкции языка программирования\n• Обеспечена кросс-платформенность за счет абстрагирования от целевой архитектуры\n• Продемонстрирована эффективность использования графа потока управления для генерации кода\n\nКомпилятор может быть использован для изучения принципов работы компиляторов и для дальнейшего развития в поддержки более сложных конструкций языка.'
        }
    ]
    
    # Helper function to add section
    def add_section(title, content=None):
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        doc.add_paragraph()
        
        if content:
            doc.add_paragraph(content)
            doc.add_paragraph()
    
    def add_code_block(code):
        code_para = doc.add_paragraph()
        code_run = code_para.add_run(code)
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(9)
        code_para.paragraph_format.left_indent = Inches(0.5)
        doc.add_paragraph()
    
    # Process content
    for section in content:
        if section['title'] in ['Цель', 'Задачи', 'Ход работы', 'Выводы']:
            add_section(section['title'], section['content'])
        
        if section.get('subsections'):
            for subsection in section['subsections']:
                if subsection['title'].startswith(('1.', '2.', '3.', '4.', '5.')):
                    subtitle_para = doc.add_paragraph()
                    subtitle_run = subtitle_para.add_run(subsection['title'])
                    subtitle_run.bold = True
                    subtitle_run.font.size = Pt(12)
                    
                    if subsection.get('content'):
                        doc.add_paragraph(subsection['content'])
                        doc.add_paragraph()
                    
                    if subsection.get('code'):
                        add_code_block(subsection['code'])
                    
                    if subsection.get('subtasks'):
                        for task in subsection['subtasks']:
                            task_para = doc.add_paragraph(task, style='List Bullet')
    
    # Add signatures
    doc.add_page_break()
    
    for _ in range(8):
        doc.add_paragraph()
    
    sig_para = doc.add_paragraph()
    sig_para.add_run('____________________ / М.С. Маршалов')
    
    for _ in range(2):
        doc.add_paragraph()
    
    sig2_para = doc.add_paragraph()
    sig2_para.add_run('____________________ / Ю.Д. Кореньков')
    
    # Save document
    doc.save('ProgLangRep3.docx')
    print('DOCX создан успешно: ProgLangRep3.docx')

if __name__ == '__main__':
    create_docx()